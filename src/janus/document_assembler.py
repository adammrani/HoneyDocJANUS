"""
src/janus/document_assembler.py
Assemble the three JANUS layers into a single python-docx Document.

  Layer 1 — Narrative (visible): framed LLM content, rendered as headings/paras.
  Layer 2 — CI1 (invisible): hidden white 1pt instruction for LLM agents + core
            metadata. Plus the Canarytoken beacon as an INCLUDEPICTURE field.
  Layer 3 — CI3 (semi-visible): an "Annexe" page of fake credentials.

The INCLUDEPICTURE field is injected via raw OXML (`OxmlElement`), because the
python-docx public API cannot create external-content fields. When Word opens
the document with external content enabled, it fetches the token_url — that
fetch is the Canarytoken trigger.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from src.janus import ci_credential_trap, ci_prompt_injection, narrative_layer
from src.core.logger import log

_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _add_hidden_run(paragraph, text: str) -> None:
    """Append a white, 1pt run so the text is effectively invisible in Word."""
    run = paragraph.add_run(text)
    run.font.color.rgb = _WHITE
    run.font.size = Pt(1)


def _inject_includepicture(doc: Document, token_url: str) -> None:
    """
    Insert a hidden INCLUDEPICTURE field pointing at token_url via raw OXML.

    Field structure: fldChar(begin) + instrText + fldChar(end), wrapped in a
    white 1pt run so it is invisible to a human reader.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # Make the carrier run invisible.
    rpr = run._element.get_or_add_rPr()
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FFFFFF")
    rpr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "2")  # half-points => 1pt
    rpr.append(sz)

    # fldChar begin
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)

    # instrText: the actual field code
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' INCLUDEPICTURE "{token_url}" \\d '
    run._element.append(instr)

    # fldChar end
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_end)


def _render_narrative(doc: Document, narrative_text: str) -> None:
    """Turn framed text into headings and paragraphs."""
    lines = narrative_text.split("\n")
    first_content_seen = False
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if not first_content_seen:
            doc.add_heading(line, level=1)
            first_content_seen = True
        elif len(line) < 60 and line.endswith(":"):
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)


def assemble_document(
    content: str,
    doc_type: str,
    token_id: str,
    token_url: str,
    enable_janus: bool = True,
    enable_ci3: bool = True,
) -> Document:
    """
    Build and return a python-docx Document with all requested layers.
    """
    doc = Document()

    # ── Layer 1: Narrative (visible) ─────────────────────
    narrative_text = narrative_layer.format_narrative(content, doc_type)
    _render_narrative(doc, narrative_text)

    # ── Canarytoken beacon (always embedded) ─────────────
    _inject_includepicture(doc, token_url)

    # ── Layer 2: CI1 (invisible) ─────────────────────────
    if enable_janus:
        hidden_para = doc.add_paragraph()
        _add_hidden_run(hidden_para, ci_prompt_injection.get_hidden_text(token_id))

        meta = ci_prompt_injection.get_metadata_injection()
        cp = doc.core_properties
        cp.subject = meta.get("subject", "")
        cp.keywords = meta.get("keywords", "")
        cp.comments = meta.get("comments", "")
        cp.category = meta.get("category", "")

    # ── Layer 3: CI3 (semi-visible) ──────────────────────
    if enable_ci3:
        doc.add_page_break()
        creds = ci_credential_trap.generate_credentials(doc_type, token_id)
        block = ci_credential_trap.format_credentials_block(creds)
        for line in block.split("\n"):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

    log.info(
        "Document assembled (type=%s janus=%s ci3=%s token=%s)",
        doc_type, enable_janus, enable_ci3, token_id,
    )
    return doc
